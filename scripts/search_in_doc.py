"""
Pour generer l'exe:

copier ce fichier dans search_in_docs_release
lancer C:\Python39\Scripts\autopytoexe.exe
ouvrir les settings et importer to_exe_config.json
lancer le truc
le resultat arrive dans output.
le deplacer depuis output.
copy output\search*.exe .
"""
import docx

import odf
import odf.opendocument

import os
import sys

strLocalPath = os.path.dirname(sys.modules[__name__].__file__)
if strLocalPath == "": strLocalPath = './'
sys.path.append(strLocalPath+"/../alex_pytools/")
import stringtools

def removeAccent(s):
    return stringtools.removeAccentString(s)

def searchInDocx(filename, s):
    """
    search s in filename.
    return nbr paragraph, nbr match
    """
    if ".docx" not in filename:
        return 0,0
    s = removeAccent(s.lower())
    doc = docx.Document(filename)
    #~ print(dir(doc))
    #~ print("DBG: nbr paragraph: %s" % len(doc.paragraphs))
    cpt = 0
    bFirstTime = 1
    allparas = doc.paragraphs
    for nNumPara,p in enumerate(allparas):
        p = p.text
        #~ print("DBG: paragraph: " + p)
        if s in removeAccent(p.lower()):
            if bFirstTime:
                bFirstTime = 0
                print("***** %s:" % filename)
            print("  paragraph %4d: %s" % ((nNumPara+1),p))
            if len(p)>120:
                print("") # add an extra line after big blocks
            cpt += 1
    return len(doc.paragraphs),cpt
    
def searchInOdt(filename, s):
    """
    search s in filename.
    return nbr paragraph, nbr match
    """
    from odf import text, teletype
    if ".odt" not in filename:
        return 0,0
    s = removeAccent(s.lower())

    doc = odf.opendocument.load(filename)
    #~ print(dir(doc))
    allparas = doc.getElementsByType(text.P)
    #~ print("DBG: nbr paragraph: %s" % len(allparas))
    cpt = 0
    bFirstTime = 1
    for nNumPara,p in enumerate(allparas):
        p = str(p)
        #~ print("DBG: paragraph: " + p)
        if s in removeAccent(p.lower()):
            if bFirstTime:
                bFirstTime = 0
                print("***** %s:" % filename)
            print("  paragraph %4d: %s" % ((nNumPara+1),p))
            if len(p)>120:
                print("") # add an extra line after big blocks
            cpt += 1
    return len(allparas),cpt
    
def extractParaFromPdf(doc):
    """
    cut text in doc in paragraph and return them
    """
    if 0:
        # text per page:
        for numpage in range(doc.pageCount):
            page = doc.load_page(numpage) #put here the page number
            page_to_text = page.get_text("text")
            print("page: %s" % page_to_text )
    
    # text per block (closer to paragraph, titles are also paragraph)
    blocks = []
    for numpage in range(doc.pageCount):
        bs = [x[4] for x in doc[numpage].get_text("blocks")]
        blocks.extend(bs)
    #~ for b in blocks:
        #~ print("block: %s" % b )
    return blocks
    
def searchInPdf(filename, s):
    """
    search s in filename.
    return nbr paragraph, nbr match
    """
    if ".pdf" not in filename:
        return 0,0
        
    import fitz

    s = removeAccent(s.lower())
    
    doc = fitz.open(filename)
    allparas = extractParaFromPdf(doc)
    cpt = 0
    bFirstTime = 1
    for nNumPara,p in enumerate(allparas):
        p = str(p)
        #~ print("DBG: paragraph: " + p)
        if s in removeAccent(p.lower()):
            if bFirstTime:
                bFirstTime = 0
                print("***** %s:" % filename)
            print("  paragraph %4d: %s" % ((nNumPara+1),p))
            cpt += 1
    return len(allparas),cpt
    
def searchInDocs(path,s,bRecurse=1,nVerbose=0):
    """
        - nVerbose: 0: rien, 1: juste quelques infos, 2: all filename treated
    """
    if path[-1] not in "/\\": path += os.sep
    listFiles = os.listdir(path)
    listFiles = sorted(listFiles)
    cptfilestot = 0
    cptparatot = 0
    cptmatchfiles = 0
    cptmatchlines = 0
    for f in listFiles:
        cfn = path + f
        if os.path.isfile(cfn) and f[0] != '~' and f[:2] != '._': # ~: lock file ._: des trucs de la corbeille ou bizarre
            name,ext = os.path.splitext(f)
            ext = ext.lower()
            if nVerbose>1: print("DBG: handling '%s'" % cfn)
            try:
                if ext in [".docx",".doc"]:
                    nbrpara,match = searchInDocx( cfn,s )
                elif ext == ".odt":
                    nbrpara,match = searchInOdt( cfn,s )
                elif ext == ".pdf":
                    nbrpara,match = searchInPdf( cfn,s )
                else:
                    if nVerbose>1: print("DBG: => skipping '%s'" % cfn)
                    continue
            except BaseException as err:
                if nVerbose>0: print("WRN: problem with file '%s', err: %s" % (cfn,err)) 
                continue
        else:
            if bRecurse: 
                if os.path.isdir(cfn):
                    nfa,npa,nfm,npm = searchInDocs(cfn,s,bRecurse=bRecurse,nVerbose=nVerbose)
                    cptfilestot += nfa
                    cptparatot += npa
                    cptmatchfiles += nfm
                    cptmatchlines += npm
            continue
        cptfilestot += 1
        cptparatot += nbrpara
        if match > 0:
            cptmatchfiles += 1
            cptmatchlines += match
            if match > 5:
                print("----- %s - end" % cfn)

    return cptfilestot, cptparatot,cptmatchfiles,cptmatchlines
    
    
#~ searchInDocx("test/1988.docx", "autoroute")
#~ searchInDocs("test", "autoroute")

if __name__ == "__main__":
    print( "\n  Search in docx/odt v1.03 by Alma\n  Searching into .pdf, doc, docx and odt")
    if len(sys.argv) < 2:
        print( "  syntax: %s <string_to_match> [1 ou 2 (for verbose/extra verbose)]\n\n  eg: %s autoroute \n" % ((sys.argv[0],)*2))
        sys.exit(-1)
    print("\n")
    
    nVerbose = 0
    if len(sys.argv)>2:
        if sys.argv[2]=='1' or sys.argv[2]=='2':
            nVerbose = int(sys.argv[2])
            
    strToMatch = sys.argv[1]
    if nVerbose>0: print("INF: nVerbose: %s" % nVerbose ) 
    nfa,npa,nfm,npm = searchInDocs(".",strToMatch,nVerbose=nVerbose)
    print("\nNbr Analysed Files: %d\nNbr Matching Files: %d\nNbr Total Paragraph Analysed  : %d\nNbr Total Paragraph With Match: %d" % (nfa,nfm,npa,npm) )
      
